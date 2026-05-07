"""
generate_report.py
Generates RENALDO_Report_v1.docx — a formatted Word document.
Run with: python generate_report.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ──
NAVY       = RGBColor(0x0F, 0x2D, 0x4A)
BLUE       = RGBColor(0x1A, 0x6F, 0xA8)
BLUE_LIGHT = RGBColor(0xE8, 0xF3, 0xFB)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT  = RGBColor(0x4A, 0x55, 0x68)
GRAY_LIGHT = RGBColor(0xF4, 0xF6, 0xF9)
GREEN      = RGBColor(0x00, 0xB0, 0x50)
ORANGE     = RGBColor(0xFF, 0x99, 0x00)
RED        = RGBColor(0xFF, 0x00, 0x00)
YELLOW_G   = RGBColor(0x92, 0xD0, 0x50)
YELLOW     = RGBColor(0xFF, 0xFF, 0x00)


def set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val",   "single"))
            el.set(qn("w:sz"),    val.get("sz",    "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para_border_left(para, hex_colour="1A6FA8", width_pt=18):
    """Add a left border (blue bar) to a paragraph — for callout boxes."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(width_pt))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_colour)
    pBdr.append(left)
    pPr.append(pBdr)


def para_shade(para, hex_colour="E8F3FB"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    pPr.append(shd)


doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default body font ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_TEXT

# ── Heading styles ──
for lvl, size, bold, colour in [
    ("Heading 1", 18, True,  NAVY),
    ("Heading 2", 13, True,  BLUE),
    ("Heading 3", 11, True,  NAVY),
]:
    s = doc.styles[lvl]
    s.font.name  = "Calibri"
    s.font.size  = Pt(size)
    s.font.bold  = bold
    s.font.color.rgb = colour
    s.paragraph_format.space_before = Pt(14)
    s.paragraph_format.space_after  = Pt(4)


# ════════════════════════════════════════════
# COVER BLOCK
# ════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_shade(title_para, "0F2D4A")
run = title_para.add_run("RENALDO")
run.font.name  = "Calibri"
run.font.size  = Pt(28)
run.font.bold  = True
run.font.color.rgb = WHITE

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_shade(sub_para, "0F2D4A")
run = sub_para.add_run("RarE kidNey dAta compLeteness DashbOard")
run.font.name  = "Calibri"
run.font.size  = Pt(13)
run.font.color.rgb = RGBColor(0xA0, 0xC4, 0xE0)

org_para = doc.add_paragraph()
org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_shade(org_para, "0F2D4A")
run = org_para.add_run("UK Kidney Association  ·  RaDaR Registry")
run.font.name  = "Calibri"
run.font.size  = Pt(10)
run.font.color.rgb = RGBColor(0xA0, 0xC4, 0xE0)

ver_para = doc.add_paragraph()
ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_shade(ver_para, "1A4A6E")
run = ver_para.add_run("Version 1.0  ·  April 2026  ·  renaldo.onrender.com")
run.font.name  = "Calibri"
run.font.size  = Pt(9)
run.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)

doc.add_paragraph()  # spacer


# ════════════════════════════════════════════
# HELPER FUNCTIONS
# ════════════════════════════════════════════
def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def h3(text): doc.add_heading(text, level=3)


def body(text, bold_parts=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    if bold_parts:
        parts = text.split("**")
        for i, part in enumerate(parts):
            r = p.add_run(part)
            if i % 2 == 1:
                r.bold = True
    else:
        p.add_run(text)
    return p


def callout(text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    para_shade(p, "E8F3FB")
    add_para_border_left(p, "1A6FA8", 18)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.bold      = True
    r.font.size      = Pt(11)
    return p


def formula(text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    para_shade(p, "F4F6F9")
    r = p.add_run(text)
    r.font.name  = "Courier New"
    r.font.size  = Pt(10)
    r.font.color.rgb = NAVY
    return p


def note(text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run("Note: ")
    r.bold = True
    r.font.color.rgb = BLUE
    r2 = p.add_run(text)
    r2.font.color.rgb = GRAY_TEXT
    r2.font.size = Pt(9.5)
    return p


def styled_table(headers, rows, header_bg="1A6FA8", alt_bg="F4F6F9"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9.5)
        r.font.name = "Calibri"

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg  = alt_bg if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            r.font.name = "Calibri"
            r.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()
    return t


# ════════════════════════════════════════════
# SECTION 1 — BACKGROUND
# ════════════════════════════════════════════
h1("1.  Background")

body("The **National Registry of Rare Kidney Diseases (RaDaR)** is a national registry collecting longitudinal clinical data on patients with rare and complex kidney conditions across the UK. It is managed by the **UK Kidney Association (UKKA)** and brings together data from renal units across England, Wales, Scotland, and Northern Ireland.", bold_parts=True)

body("RaDaR currently holds records for over 39,000 patients across 33 disease cohort groups. As with any large clinical registry, the quality and completeness of data varies across sites and cohorts. Poor completeness limits the scientific and clinical value of the registry.")

body("**RENALDO** (RarE kidNey dAta compLeteness DashbOard) was developed to provide a clear, regularly updated view of data completeness across the entire RaDaR dataset — enabling data managers, cohort leads, and the UKKA to quickly identify where data collection requires improvement.", bold_parts=True)


# ════════════════════════════════════════════
# SECTION 2 — PATIENT POPULATION
# ════════════════════════════════════════════
h1("2.  Patient Population")
h2("2.1  Inclusion Criteria")

body("Only patients meeting all of the following criteria are included in the analysis:")

styled_table(
    ["Criterion", "Rule"],
    [
        ["Source", "Demographics record must have source_type = 'RADAR'"],
        ["Patient type", "test = FALSE and control = FALSE"],
        ["Group type", "Must be enrolled in a group of type COHORT"],
        ["Excluded cohorts", "Must not belong to any excluded group IDs (see Section 2.2)"],
    ]
)

h2("2.2  Excluded Cohorts")
body("The following cohort groups are excluded from all calculations. These are either non-standard groups, administrative groups, or cohorts not part of the core RaDaR rare kidney disease dataset:")

styled_table(
    ["Group ID", "Name / Reason for Exclusion"],
    [
        ["137", "NURTuRE-CKD"],
        ["149", "NephroS"],
        ["152", "NaHUS"],
        ["18, 19", "Administrative / legacy groups"],
        ["161, 174, 182, 184, 194, 140, 220, 222", "Withdrawn consent / non-standard groups"],
    ]
)

h2("2.3  Total Patient Count")
body("After applying all inclusion and exclusion criteria, the total analysed population is:")
callout("N = 39,178 patients")
body("This is the denominator used for all overall (Section A — Overall RaDaR) completeness calculations.")


# ════════════════════════════════════════════
# SECTION 3 — VARIABLES
# ════════════════════════════════════════════
h1("3.  Variables Assessed")
h2("3.1  Demographic Variables")
body("The following 11 demographic variables are assessed for all patients across both the overall RaDaR section and each of the 33 cohort sections:")

styled_table(
    ["ID", "Variable", "Required", "Source"],
    [
        ["A.1",  "First Name",      "Yes", "patient_demographics.first_name"],
        ["A.2",  "Last Name",       "Yes", "patient_demographics.last_name"],
        ["A.3",  "Date of Birth",   "Yes", "patient_demographics.date_of_birth"],
        ["A.4",  "Date of Death",   "No",  "patient_demographics.date_of_death"],
        ["A.5",  "Cause of Death",  "No",  "patient_demographics.cause_of_death"],
        ["A.6",  "Gender",          "Yes", "patient_demographics.gender"],
        ["A.7",  "Ethnicity",       "Yes", "patient_demographics.ethnicity_id"],
        ["A.8",  "Nationality",     "Yes", "patient_demographics.nationality_id"],
        ["A.9",  "Email Address",   "No",  "patient_demographics.email_address"],
        ["A.10", "Diagnosis",       "Yes", "patient_diagnoses table (any record present)"],
        ["A.11", "NHS Number",      "Yes", "patient_numbers table (any record present)"],
    ]
)
body("Variables marked Required are part of the RaDaR minimum dataset and should be present for every registered patient.")

h2("3.2  Clinical Summary Rows")
body("These rows appear in each section but show counts rather than % missing. They are not included in the section completeness score calculation:")

styled_table(
    ["Row", "What it shows"],
    [
        ["TOTAL_PATIENTS",      "Total patients enrolled in that cohort"],
        ["ADULTS",              "Patients aged ≥ 18 at time of analysis"],
        ["CHILDREN",            "Patients aged < 18 at time of analysis"],
        ["KIDNEY_FAILURE",      "Patients with evidence of kidney failure (see Section 5)"],
        ["TRANSPLANT_SINGLE",   "Patients with exactly one transplant"],
        ["TRANSPLANT_MULTIPLE", "Patients with two or more transplants"],
        ["FOLLOW_UP",           "Median follow-up duration in years (see Section 6)"],
    ]
)


# ════════════════════════════════════════════
# SECTION 4 — COMPLETENESS CALCULATIONS
# ════════════════════════════════════════════
h1("4.  Completeness Calculations")
h2("4.1  Variable-Level % Missing")

body("For each variable and each cohort, the percentage of missing values is calculated as:")
formula("% Missing  =  (Patients with no recorded value  ÷  Total patients in cohort)  ×  100")

body("A value is considered missing if it is:")
for item in [
    "NULL in the database, or",
    "An empty string after stripping whitespace",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10.5)

doc.add_paragraph()
h3("Special Cases")

body("**Cause of Death** — denominator is restricted to deceased patients only, as it is not meaningful to measure this field for living patients:", bold_parts=True)
formula("% Missing (Cause of Death)  =  Deceased patients with no cause of death  ÷  Total deceased patients  ×  100")

body("**Email Address** — a list of 46 known placeholder and default email addresses is maintained (e.g. radar@radar.org, noemailaddress@radar.radar). Any patient whose email matches one of these placeholders is treated as missing, even if a value technically exists in the database:", bold_parts=True)
formula("Missing email  =  NULL emails  +  emails matching any placeholder in the 46-entry list")

body("**NHS Number** — completeness is checked by the presence of any record for that patient in the patient_numbers table, not the demographics row itself.", bold_parts=True)

body("**Diagnosis** — completeness is checked by the presence of any record for that patient in the patient_diagnoses table.", bold_parts=True)

h2("4.2  Colour Coding")
body("Each variable row is colour-coded based on its % missing value:")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0]
for i, h in enumerate(["Colour", "% Missing", "Interpretation"]):
    set_cell_bg(hdr.cells[i], "1A6FA8")
    p = hdr.cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9.5)

colour_rows = [
    ("00B050", "0 – 20%",   "Good completeness",    "FFFFFF"),
    ("92D050", "20 – 40%",  "Acceptable",            "1A4A00"),
    ("FFFF00", "40 – 60%",  "Needs attention",       "4A4A00"),
    ("FF9900", "60 – 80%",  "Poor",                  "4A2500"),
    ("FF0000", "80 – 100%", "Critical",              "FFFFFF"),
]
for bg, rng, interp, fg in colour_rows:
    row = t.add_row()
    set_cell_bg(row.cells[0], bg)
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run("  ")
    r0.font.size = Pt(9.5)

    set_cell_bg(row.cells[1], "FFFFFF")
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(rng)
    r1.font.size = Pt(9.5); r1.font.name = "Calibri"

    set_cell_bg(row.cells[2], "FFFFFF")
    p2 = row.cells[2].paragraphs[0]
    r2 = p2.add_run(interp)
    r2.font.size = Pt(9.5); r2.font.name = "Calibri"

doc.add_paragraph()

h2("4.3  Section-Level % Complete")
body("Each section header in the dashboard displays an overall completeness percentage, calculated as:")
formula(
    "Section % Complete  =  100  −  (Sum of % Missing for all n variables  ÷  n)\n\n"
    "Where n = number of variables with a % missing value\n"
    "(TOTAL_PATIENTS, FOLLOW_UP, KF and Transplant rows are excluded)"
)

body("Example: a cohort with 11 demographic variables having % missing values of 0, 2, 5, 10, 15, 30, 45, 60, 70, 80, 90:")
formula(
    "Average % Missing  =  (0+2+5+10+15+30+45+60+70+80+90) ÷ 11  =  407 ÷ 11  ≈  37%\n\n"
    "Section % Complete  =  100 − 37  =  63%"
)
note("This is an unweighted average — every variable contributes equally regardless of how many patients it applies to. It is a quick visual indicator, not a weighted audit score.")


# ════════════════════════════════════════════
# SECTION 5 — KIDNEY FAILURE
# ════════════════════════════════════════════
h1("5.  Kidney Failure Definition")

body("A patient is classified as having Kidney Failure (KF) if they meet any of the following criteria, based on data recorded in RaDaR:")

styled_table(
    ["Criterion", "Definition"],
    [
        ["Transplant", "Any record present in the transplants table"],
        ["Dialysis",   "Any record present in the dialysis table"],
        ["eGFR < 15",  "Two or more eGFR readings < 15 ml/min/1.73m², recorded ≥ 28 days apart, with no eGFR ≥ 15 recorded between the first and last low reading"],
    ]
)

body("The eGFR criterion uses observation ID 47 in the results table.")

p = doc.add_paragraph(style="Normal")
para_shade(p, "FFF3CD")
add_para_border_left(p, "FF9900", 18)
r = p.add_run("Important: KF classification is based solely on RaDaR data and is not linked to the UK Renal Registry (UKRR). Patients on dialysis or with a transplant managed entirely outside RaDaR-linked units will not be captured.")
r.font.size = Pt(10); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x4A, 0x25, 0x00)
doc.add_paragraph()

h2("5.1  Transplant Counting")
body("Transplants are counted per patient using COUNT(DISTINCT date) from the transplants table — each unique transplant date counts as one transplant. Patients are then split into:")
for item in [
    "Single transplant: exactly 1 distinct transplant date",
    "Multiple transplants: 2 or more distinct transplant dates",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10.5)
doc.add_paragraph()


# ════════════════════════════════════════════
# SECTION 6 — FOLLOW-UP
# ════════════════════════════════════════════
h1("6.  Follow-Up Calculation")
body("Follow-up measures the duration of a patient's active monitoring in RaDaR, expressed in years.")

h2("6.1  Formula")
formula(
    "Follow-up (years)  =  (End Date  −  Enrolment Date)  ÷  365.25\n\n"
    "Enrolment Date  =  group_patients.from_date  (the 'Recruited On' date in RaDaR)\n"
    "                   Fallback: earliest cohort membership date across all cohorts\n\n"
    "End Date hierarchy:\n"
    "  1.  Date of death          (if the patient is deceased)\n"
    "  2.  Last activity date     (latest date in results or medications tables)\n"
    "  3.  Today's date           (if no activity has ever been recorded)"
)

body("The dashboard reports the median follow-up and interquartile range (IQR) for each section.")

h2("6.2  Special Cases")
styled_table(
    ["Situation", "How it is handled"],
    [
        ["Last activity is before enrolment date",  "Treated as a historical batch upload — end date set to today"],
        ["Date of death is before enrolment date",  "Data entry error — patient excluded from follow-up calculation"],
        ["No enrolment date available",             "Cannot calculate — patient excluded from median"],
    ]
)

h2("6.3  Known Data Quality Issue")

p = doc.add_paragraph(style="Normal")
para_shade(p, "FFE7E7")
add_para_border_left(p, "FF0000", 18)
r = p.add_run("Across the full RaDaR dataset, approximately 18 patients have a date of death recorded earlier than their enrolment date. This is a data entry error at the enrolling unit. These patients are excluded from the follow-up calculation. Correcting the dates in the RaDaR front end would restore them to the analysis automatically on the next data refresh.")
r.font.size = Pt(10); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x6B, 0x00, 0x00)
doc.add_paragraph()


# ════════════════════════════════════════════
# SECTION 7 — ARCHITECTURE
# ════════════════════════════════════════════
h1("7.  Dashboard Architecture")
h2("7.1  Data Pipeline")
body("The dashboard operates on a static JSON model — there is no live database connection on the hosted version:")

steps = [
    ("Step 1", "run_all.py runs locally — opens SSH tunnel to RaDaR PostgreSQL database, runs all queries (demographics, cohort counts, follow-up, KF, transplants), writes output/completeness.json"),
    ("Step 2", "completeness.json is committed to GitHub"),
    ("Step 3", "Render.com auto-deploys from GitHub — dashboard reads completeness.json at startup. No database credentials required on the server."),
]
for label, text in steps:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True; r1.font.color.rgb = BLUE
    r2 = p.add_run(text)
    r2.font.color.rgb = GRAY_TEXT
doc.add_paragraph()

h2("7.2  Update Frequency")
body("The dashboard is updated manually by running the analytics pipeline locally and pushing the resulting JSON file to GitHub. Render.com redeploys automatically within approximately 2 minutes of each push.")

h2("7.3  Technology Stack")
styled_table(
    ["Component", "Technology"],
    [
        ["Dashboard framework", "Plotly Dash (Python)"],
        ["UI components",       "Dash Bootstrap Components (Bootstrap 5)"],
        ["Data processing",     "pandas"],
        ["Database connection", "PostgreSQL via SSH tunnel (psycopg2 + sshtunnel)"],
        ["Hosting",             "Render.com"],
        ["Version control",     "GitHub — encrypted000/RENALDO"],
    ]
)


# ════════════════════════════════════════════
# SECTION 8 — LIMITATIONS
# ════════════════════════════════════════════
h1("8.  Limitations")

limitations = [
    ("Static data", "The dashboard reflects the state of RaDaR at the time the analytics pipeline was last run. It does not update in real time."),
    ("RaDaR-only KF classification", "Kidney failure status is derived from RaDaR records only. Patients receiving renal replacement therapy outside RaDaR-linked pathways will not be captured."),
    ("Unweighted completeness score", "The section % complete is a simple average across variables and does not weight by variable importance or patient count."),
    ("Email placeholder list", "The list of 46 placeholder emails is maintained manually. New placeholder addresses added to RaDaR after the last update of this list will not be treated as missing."),
    ("Follow-up based on last activity", "Follow-up end date is the latest entry in results or medications. Patients who are alive and enrolled but have had no activity recorded will appear to have follow-up ending today, which may overestimate their active monitoring period."),
]

for i, (title, text) in enumerate(limitations, 1):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f"{i}.  {title} — ")
    r1.bold = True; r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.color.rgb = GRAY_TEXT


# ════════════════════════════════════════════
# SECTION 9 — GLOSSARY
# ════════════════════════════════════════════
h1("9.  Glossary")
styled_table(
    ["Term", "Definition"],
    [
        ["RaDaR",          "National Registry of Rare Kidney Diseases"],
        ["UKKA",           "UK Kidney Association"],
        ["RENALDO",        "RarE kidNey dAta compLeteness DashbOard"],
        ["Completeness",   "The proportion of expected values that have been recorded"],
        ["% Missing",      "Patients with no recorded value ÷ total patients × 100"],
        ["KF",             "Kidney Failure"],
        ["eGFR",           "Estimated Glomerular Filtration Rate — a measure of kidney function (ml/min/1.73m²)"],
        ["IQR",            "Interquartile Range — the range between the 25th and 75th percentile"],
        ["SSH Tunnel",     "Secure encrypted connection used to access the RaDaR database"],
        ["source_type RADAR", "Records entered directly into RaDaR (as opposed to imported from other systems)"],
    ]
)


# ════════════════════════════════════════════
# FOOTER PARAGRAPH
# ════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph(style="Normal")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Report prepared by the UK Kidney Association data team  ·  Dashboard: renaldo.onrender.com  ·  Source: github.com/encrypted000/RENALDO")
r.font.size = Pt(8.5)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


# ════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════
out = "RENALDO_Report_v1.docx"
doc.save(out)
print(f"Saved: {out}")
